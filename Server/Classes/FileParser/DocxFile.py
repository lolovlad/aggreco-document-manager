from pathlib import Path
from re import findall

from docx import Document
from docx.table import Table

from ..Models.Fileschame import *

from .Cell import Cell
from .Table import Table as CustomTable


from .File import File


class DocxFile(File):
    def __init__(self, path: Path):
        self.__path: Path = path
        self.__file: Document = Document(path)

    def get_all_parser_table_in_file(self, re: str) -> list[tuple[str, CustomTable]]:
        tables = self.__file.tables
        id_custom_table = []
        for i, table in enumerate(tables):
            if self.is_table_schema(table, re):
                custom_table = self.__create_table(table)
                key_table = self.__get_key_table(custom_table[0][0])
                if key_table is not None:
                    table = self.delete_system_row(i)
                    id_custom_table.append((key_table, i))

        self.__file.save(self.__path)
        self.__file = Document(self.__path)
        tables = self.__file.tables

        custom_tables = []
        for key_table, i in id_custom_table:
            custom_table = self.__create_table(tables[i])
            custom_tables.append((key_table, custom_table))
        return custom_tables

    def render(self, schemas: dict):
        pass

    def __create_table(self, table: Table) -> CustomTable:
        custom_table = CustomTable(len(table.rows), len(table.columns))
        cells = self.get_list_cells(table)
        cells = self.__create_list_cells(list(cells))
        for cell in cells:
            custom_table.add_cell(cell)
        return custom_table

    def is_table_schema(self, table: Table, re: str) -> bool:
        for k in table._cells:
            if len(findall(re, k.text)) > 0:
                return True
        return False

    def get_list_cells(self, table: Table):
        cells = set()
        for y, row in enumerate(table.columns):
            for x, cell in enumerate(row.cells):
                cell_loc = self.__get_coord_cell(cell)
                cells.add((cell.text, *cell_loc))
        return cells

    def delete_system_row(self, id_table: int) -> Table:
        self.__file.tables[id_table]._tbl.remove(self.__file.tables[id_table].rows[0]._tr)
        return self.__file.tables[id_table]

    def __get_key_table(self, cell) -> str | None:
        key_str = "".join(cell.text.split()).lower()
        main = key_str.split("_")
        if len(main) > 1:
            if main[1] == "table":
                return key_str
        return None

    def __get_coord_cell(self, cell):
        tc = cell._tc
        try:
            return (tc.top, tc.bottom, tc.left, tc.right)
        except ValueError:

            return (tc.top, self.__get__bottom_coord_cell(cell._parent, cell, (tc.top, tc.left)), tc.left, tc.right)

    def __get__bottom_coord_cell(self, table: Table, cell, coord) -> int:
        text_cell = cell.text
        bottom_coord = coord[0]
        next_text_cell = table.cell(bottom_coord, coord[1])
        while text_cell == next_text_cell.text:
            bottom_coord += 1
            next_text_cell = table.cell(bottom_coord, coord[1])
        return bottom_coord

    def __create_list_cells(self, cells: list) -> list[Cell]:
        return [Cell(top=cell[1], bottom=cell[2], left=cell[3], right=cell[4], text=cell[0]) for cell in cells]

    def __create_list_cells_schames(self, cells: list[Cell], re: str):
        cells_schemas = [cell.get_schemas() for cell in cells]
        return cells_schemas